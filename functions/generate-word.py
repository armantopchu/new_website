# functions/generate_word.py
from docx import Document
from flask import Flask, jsonify, request
import os

app = Flask(__name__)

@app.route('/.netlify/functions/generate-word', methods=['POST'])
def generate_word():
    data = request.get_json()

    # Implement your Word file generation logic here
    # Example: Create a Word document with the received data
    doc = Document()
    doc.add_paragraph(f"First Name: {data['first_name']}")
    doc.add_paragraph(f"Surname: {data['surname']}")
    doc.save('/tmp/test_chatGPT10.docx')

    return jsonify({'success': True, 'message': 'Word file generated successfully'})