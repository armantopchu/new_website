from flask import Flask, render_template, request, jsonify
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.oxml.shared import qn  # Add this line to import qn
import os

app = Flask(__name__)

@app.route('/')
def home():
    return render_template('index.html')

@app.route('/submit-form', methods=['POST'])
def submit_form():
    # Retrieve form data from the request
    data = request.form

    # Call your function to generate Word file using the form data
    generate_word_file(data)

    # Return a response (for simplicity, just echoing back the received data)
    return jsonify({'success': True, 'data': data})

def set_font_and_spacing(paragraph, font_name, font_size, spacing):
    run = paragraph.add_run('')  
    font = run.font
    font.name = font_name
    font.size = Pt(font_size)
    paragraph.space_after = Pt(spacing)

def add_heading(document, text, spacing_before, spacing_after):
    heading = document.add_paragraph(text, style='Heading 1')
    run = heading.runs[0]
    font = run.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0, 0, 0)  # Standard black color
    heading.paragraph_format.space_before = Pt(spacing_before)  # Set paragraph spacing before
    heading.paragraph_format.space_after = Pt(spacing_after)  # Set paragraph spacing after

def insert_image_right_aligned(document, image_path, width):
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(image_path, width=Inches(width))

    # Align the image to the right
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

def add_address(document, lines, alignment):
    for line in lines:
        paragraph = document.add_paragraph(line)
        paragraph.alignment = alignment
        paragraph.style.font.name = 'Calibri'  # Set Calibri as the default font
        paragraph.style.paragraph_format.space_after = Pt(0)  # No spacing after paragraphs of the same style

def add_table(document, data, is_bold_total=False):
    table_data = data

    table = document.add_table(rows=len(table_data), cols=len(table_data[0]))
    
    # Set table cell formatting
    for i, row in enumerate(table_data):
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            cell.text = value

            # Set font color and background color
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(0, 0, 0)  # Set default font color to black
                
                # Set background color for all cells using graphic fill
                cell._element.get_or_add_tcPr().append(OxmlElement('w:shd'))
                cell._element.tcPr.solidFill = '000000'  # Black background color using graphic fill

            # Set background color for the first row (headers)
            if i == 0:
                cell._element.get_or_add_tcPr().append(OxmlElement('w:shd'))
                cell._element.tcPr.solidFill = '000000'  # Black background color for headers

                # Set font color for the top row (headers) to white
                for run in table.rows[0].cells[j].paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)

    # Bold the Total row if specified
    if is_bold_total:
        for run in table.rows[-1].cells[0].paragraphs[0].runs:
            run.font.bold = True

        # Bold the total amount cell
        for run in table.rows[-1].cells[1].paragraphs[0].runs:
            run.font.bold = True

    return table

def add_recommendation_paragraph(document, data):
    # Add space (empty row) before the recommendation paragraph
    document.add_paragraph()

    # Add main text with spacing between lines
    main_text = "I recommend that you invest in X,Y,Z as demonstrated below:"
    document.add_paragraph(main_text).paragraph_format.space_after = Pt(6)  # Set paragraph spacing to 6

    # Add a new table under the recommendation paragraph
    add_table(document, data)

def calculate_fees(isa_amount, fees_percent):
    fees_calc = fees_percent * isa_amount
    return fees_calc

def add_page_number_footer(document):
    section = document.sections[0]
    footer = section.footer

    paragraph = footer.paragraphs[0].clear()
    run = paragraph.add_run()

    # Add page number field
    field_code = 'PAGE'
    field = OxmlElement('w:fldSimple')
    field.set(qn('w:instr'), f' PAGE ')
    run._element.append(field)

    # Set font size for the page number
    run.font.size = Pt(10)

    # Align the paragraph to the right
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

def add_header(document, image_path, width):
    section = document.sections[0]
    header = section.header

    paragraph = header.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(image_path, width=Inches(width))  # Set the desired image size

    # Align the paragraph to the right
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

def generate_word_file(data):
    # Create a new Word document
    doc = Document()
    # Set default font and spacing
    set_font_and_spacing(doc.add_paragraph(), 'Calibri', 11, 0)  # Removed spacing for addresses

    # Add logo image aligned to the right in the header
    add_header(doc, 'logo 1.jpg', 0.7)

    # Add "From" address right-aligned
    add_address(doc, ["From", "Arman Topchu", "Apartment 14, Mary's Court,", "4 Palgrave Gardens, London", "NW1 6EW","London"],
                WD_PARAGRAPH_ALIGNMENT.RIGHT)

    # Add "To" address left-aligned
    add_address(doc, ["To:", f"Mr {data['surname']}", "Address 1", "Address 2", "Postcode", "United Kingdom"],
                WD_PARAGRAPH_ALIGNMENT.LEFT)

    # Add subject heading centered and in bold with adjusted paragraph spacing
    add_heading(doc, "RE: ISA Top-up 2023/24", 12, 12)

    # Add main text with spacing between lines
    main_text = f"Dear Mr {data['surname']},\n\nI hope you are well.\n\nIt was a pleasure to see you yesterday. " \
                "It was great to discuss X,Y,Z. In this document I plan to...\n\nPlease find an updated summary of your assets below"
    doc.add_paragraph(main_text).paragraph_format.space_after = Pt(6)  # Set paragraph spacing to 6

    # Add the first table with specific formatting
    table_data = [['Description', 'Amount', 'Commentary'],
                  ['ISA', '£25000', 'Text1'],
                  ['GIA', '£44000', 'Text2'],
                  ['SIPP', '£25000', 'Text3']]

    # Calculate the total of the second column (Amount)
    total_amount = sum(float(row[1][1:]) for row in table_data[1:])
    table_data.append(['Total', f'£{total_amount:.0f}', ''])

    # Add the updated table with bold Total row
    first_table = add_table(doc, table_data, is_bold_total=True)

    # Calculate fees
    fees_percent = 0.01  # 1%
    isa_amount = float(first_table.cell(1, 1).text[1:])  # Extract the amount from the first table
    fees_calc = calculate_fees(isa_amount, fees_percent)

    # Add the recommendation paragraph with a new table
    add_recommendation_paragraph(doc, [['Description', 'Amount', 'Commentary'],
                                       ['ISA', '£20000', 'Placeholder']])

    # Add the additional sentence about fees with space before
    doc.add_paragraph().paragraph_format.space_before = Pt(6)  # Add space before the paragraph
    fees_sentence = f"Your annual fees are {fees_percent * 100}%, which is equal to £{fees_calc:.2f}"
    doc.add_paragraph(fees_sentence).paragraph_format.space_after = Pt(6)  # Set paragraph spacing to 6

    # Add closing text with space before
    doc.add_paragraph().paragraph_format.space_before = Pt(6)  # Add space before the paragraph
    doc.add_paragraph("Thank you for your time.\n\nKind regards,\n\nArman Topchu")

    # Add page number in footer
    add_page_number_footer(doc)

    # Get the absolute path to the directory where this script is located
    script_dir = os.path.dirname(os.path.abspath(__file__))
    docx_filename = os.path.join(script_dir, 'test_chatGPT10.docx')

    # Save the Word document
    doc.save(docx_filename)
    
if __name__ == '__main__':
    app.run(debug=True, host='127.0.0.1', port=5000)